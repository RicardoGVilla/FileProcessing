import tkinter as tk
from tkinter import filedialog
import docx
import PyPDF2

# Initialize variables
file_path = None
export_instructions_selected = None
producer_input = None
product_input = None 
export_info = {}


# Function to upload files
def upload_file():
    global file_path, export_instructions_selected, producer_input, product_input, export_instructions

    if file_path is None and export_instructions_selected is None:
        file_path = filedialog.askopenfilename(filetypes=[("PDF Files", "*.pdf"), ("Word Files", "*.docx")])
        process_document(file_path)
    elif file_path and export_instructions_selected is None:
        export_instructions_selected = get_export_instructions()
        

# Function to get export instructions
def get_export_instructions():
    global export_instructions_selected, export_info
    export_instructions_selected = filedialog.askopenfilename(filetypes=[("Word Files", "*.docx")])
    process_export_instructions(export_instructions_selected)


def process_document(file):
    global file_path
    pdf_doc = PyPDF2.PdfReader(file)
    counter = 0 
    file_path = {}
    key = None
    value_buffer = ''
    for page_num in range(len(pdf_doc.pages)):
        page = pdf_doc.pages[page_num]
        for line in page.extract_text().split('\n'):
            if ":" in line:
                if key: 
                    file_path[key] = ' '.join(value_buffer.split()) 
                    value_buffer = ''  
                key, value = line.split(":", 1)
                key = key.strip()
                value_buffer = value.strip()  
                counter += 1
            else:
                value_buffer = value_buffer + ' ' + line.strip()
    if key:
        file_path[key] = ' '.join(value_buffer.split())  



def process_export_instructions(instruction_text):
    global export_info, file_path
    word_doc = docx.Document(instruction_text)

    key_mapping = {
        "Product": "product",
        "Producer": "producer",
        "Importer": "importer",
        "Lot Number": "lot"
    }
    counter = 0 
    for paragraph in word_doc.paragraphs:
        for line in paragraph.text.split('\n'): 
            counter = counter + 1
            if ":" in line:
                key, value = line.split(":", 1)
                key = key.strip()
                if key in key_mapping:
                    export_info[key_mapping[key]] = value.strip()

    return export_info



# Function to print PDF content
# def print_pdf(file_path):
#     pdf_reader = PyPDF2.PdfReader(file_path)
#     for page_num in range(len(pdf_reader.pages)):
#         page = pdf_reader.pages[page_num]
#         for line in page.extract_text().split('\n'):
#             if line.strip():
#                 print(line)

# Function to print Word content
def print_word(file_path):
    word_read = docx.Document(file_path)
    for paragraph in word_read.paragraphs:
        text = paragraph.text.strip()
        if text:
            print(text)

# Function to compare documents
def compare_documents():
    global export_info, file_path

    if export_info:
        print(export_info)
        print(file_path)
        
        for export_key, export_value in export_info.items():
            match_found = False  # Reset match_found for each export key
            for file_key, file_value in file_path.items():
                # Normalize whitespace for comparison
                normalized_export_value = ' '.join(export_value.split())
                normalized_file_value = ' '.join(file_value.split())
                
                if export_key.lower() in file_key.lower():
                    match_found = True
                    if normalized_export_value == normalized_file_value:
                        print(f"Key '{export_key}' matches with value '{export_value}' in both documents.")
                    else:
                        print(f"Key '{export_key}' matches but values do not match: Export Info: '{export_value}', File Path: '{file_value}'")
                    break  
            
            if not match_found:
                print(f"No matching key found for '{export_key}' between export label and export instructions.")
                
        if not match_found:
            print("No matching key found between export label and export instructions.")
    else: 
        response = tk.messagebox.askyesno("No Export Instructions", "No export instructions found. Do you want to upload export instructions?")
        if response:
            get_export_instructions()


def normalize_text(text):
    text = text.replace(" - ", "-")
    text = ' '.join(text.split())
    return text



# Create the main window with tkinter
root = tk.Tk()
root.title("Label Approval Tool")

# Function to select producer
def select_producer(supplier):
    global producer_input
    producer_input = supplier
    print("now you have a producer:", producer_input)

# Function to select product
def select_product(item):
    global product_input
    product_input = item 
    print("now you have an item:", product_input)

# Function to enable the export instructions button
def enable_export_instructions_button():
    global export_instructions_button
    export_instructions_button.config(state=tk.NORMAL)


# Create UI elements to upload and compare documents 
upload_button = tk.Button(root, text="Upload Document", command=upload_file)
export_instructions_button = tk.Button(root, text="Upload Export Instructions", command=get_export_instructions, state=tk.DISABLED)
producers = ['Producer 1', 'Producer 2', 'Producer 3']  
products = ['Product A', 'Product B', 'Product C']  
export_instructions = {}
producer_var = tk.StringVar(root)
producer_var.set(producers[0])
product_var = tk.StringVar(root)
product_var.set(products[0])
producer_dropdown = tk.OptionMenu(root, producer_var, *producers, command=select_producer)
product_dropdown = tk.OptionMenu(root, product_var, *products, command=select_product)
compare_button = tk.Button(root, text="Compare Documents", command=lambda: compare_documents())

# Place UI elements using grid
upload_button.grid(row=0, column=0)
export_instructions_button.grid(row=1, column=0)
producer_dropdown.grid(row=2, column=0)
product_dropdown.grid(row=2, column=1)
compare_button.grid(row=3, column=0)

# Start the Tkinter main loop
root.mainloop()